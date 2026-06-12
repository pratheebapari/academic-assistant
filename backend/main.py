from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import fitz
import os
import json
import httpx
import docx
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
HISTORY_FILE = "chat_history.json"
document_store = {"text": ""}

def load_history():
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, "r") as f:
            return json.load(f)
    return []

def save_history(history):
    with open(HISTORY_FILE, "w") as f:
        json.dump(history, f)

class QuestionRequest(BaseModel):
    question: str
    chat_id: str = "default"

@app.get("/")
def root():
    return {"message": "Academic Assistant API is running"}

@app.get("/history")
def get_history():
    return load_history()

@app.delete("/history/{chat_id}")
def delete_chat(chat_id: str):
    history = load_history()
    history = [h for h in history if h.get("chat_id") != chat_id]
    save_history(history)
    return {"message": "Chat deleted"}

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    contents = await file.read()
    text = ""
    filename = file.filename.lower()
    if filename.endswith(".pdf"):
        pdf = fitz.open(stream=contents, filetype="pdf")
        for page in pdf:
            text += page.get_text()
        pages = len(pdf)
    elif filename.endswith(".docx"):
        import io
        doc = docx.Document(io.BytesIO(contents))
        for para in doc.paragraphs:
            text += para.text + "\n"
        pages = len(doc.paragraphs)
    elif filename.endswith(".txt"):
        text = contents.decode("utf-8")
        pages = 1
    else:
        return {"message": "Unsupported file type", "pages": 0}
    document_store["text"] = text
    return {"message": f"{file.filename} uploaded successfully", "pages": pages}

@app.post("/ask")
async def ask_question(request: QuestionRequest):
    doc_text = document_store["text"]
    if doc_text:
        prompt = f"Based on this document:\n\n{doc_text[:3000]}\n\nAnswer this question clearly: {request.question}"
    else:
        prompt = request.question

    async def stream():
        full_answer = ""
        async with httpx.AsyncClient(timeout=60) as client:
            async with client.stream(
                "POST",
                "https://api.groq.com/openai/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {GROQ_API_KEY}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": "llama-3.3-70b-versatile",
                    "messages": [{
                                    "role": "system",
                                    "content": "You are an Academic Assistant — an AI built specifically to help students learn. Your purpose is to explain concepts clearly, break down complex topics into simple language, answer academic questions accurately, summarize study materials, and help students understand their notes. Always give structured, easy-to-understand answers. Use examples where helpful. If asked who you are, say: 'I am your Academic Assistant, here to help you learn, understand concepts, and make sense of your study materials.' Never say you are made by Meta or any company."
                                },
                                {"role": "user", "content": prompt}
],
                    "stream": True
                }
            ) as response:
                async for line in response.aiter_lines():
                    if line.startswith("data: "):
                        data = line[6:]
                        if data == "[DONE]":
                            break
                        try:
                            chunk = json.loads(data)
                            token = chunk["choices"][0]["delta"].get("content", "")
                            if token:
                                full_answer += token
                                yield token
                        except:
                            pass

        history = load_history()
        chat_id = request.chat_id
        existing = next((h for h in history if h.get("chat_id") == chat_id), None)
        if existing:
            existing["messages"].append({"role": "user", "text": request.question})
            existing["messages"].append({"role": "assistant", "text": full_answer})
            existing["time"] = datetime.now().strftime("%d %b, %H:%M")
        else:
            history.append({
                "chat_id": chat_id,
                "title": request.question[:40],
                "time": datetime.now().strftime("%d %b, %H:%M"),
                "messages": [
                    {"role": "user", "text": request.question},
                    {"role": "assistant", "text": full_answer}
                ]
            })
        save_history(history)

    return StreamingResponse(stream(), media_type="text/plain")